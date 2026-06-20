import os
import re
import sys
from pathlib import Path
from docx import Document
from modules.helpers import print_lg, critical_error_log

# Import formatting functions from existing builder tools
try:
    sys.path.insert(0, str(Path(__file__).resolve().parents[2]))
    from tools.build_tailored_resumes import (
        set_doc_defaults,
        add_header,
        add_section,
        add_role,
        add_bullet,
        add_inline_kv,
        add_education,
        add_publications,
        OUT_DIR
    )
except Exception as e:
    print(f"Error importing from build_tailored_resumes: {e}")
    # Fallback layout definitions in case import fails
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    OUT_DIR = Path(__file__).resolve().parents[2] / "tailored_resumes" / "docx"
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    BLUE = RGBColor(31, 77, 120)
    DARK = RGBColor(20, 20, 20)
    MUTED = RGBColor(90, 90, 90)

    def set_doc_defaults(doc: Document):
        section = doc.sections[0]
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.62)
        section.right_margin = Inches(0.62)
        styles = doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(9.6)
        normal.paragraph_format.space_after = Pt(2.5)
        normal.paragraph_format.line_spacing = 1.06

    def add_header(doc, name, subtitle):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(name).font.size = Pt(18)
        p.runs[0].bold = True
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run("yulunwu4@illinois.edu | 217-819-1856 | Champaign, IL")
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.add_run(subtitle).bold = True

    def add_section(doc, title):
        doc.add_paragraph(title.upper())

    def add_role(doc, title, org=None, date=None, location=None):
        doc.add_paragraph(f"{title} | {org or ''} | {date or ''}")

    def add_bullet(doc, text):
        doc.add_paragraph(f"• {text}")

    def add_inline_kv(doc, label, value):
        p = doc.add_paragraph()
        p.add_run(label + ": ").bold = True
        p.add_run(value)

    def add_education(doc):
        add_section(doc, "Education")
        add_role(doc, "Ph.D. in Engineering", "University of Illinois Urbana-Champaign", "Expected 2027")

    def add_publications(doc, items):
        add_section(doc, "Selected Publications")
        for item in items:
            add_bullet(doc, item)

# Global project database for Yulun Wu
PROJECTS_DATABASE = {
    "surrogate_modeling": {
        "title": "Physics-Guided ML Surrogate Modeling Researcher",
        "org": "UIUC",
        "date": "2023-Present",
        "default_bullets": [
            "Developed physics-guided CNN and multi-task surrogate models to predict both scalar strain energy and full-field displacement maps from heterogeneous microstructure inputs.",
            "Designed and implemented PCIConv/BLINKConv-style physical-condition-aware convolution modules to inject loading and boundary conditions directly into material-response surrogate models.",
            "Built unbalanced-data experiments where abundant scalar-response data supported scarce full-field response prediction, targeting reduced dependence on expensive high-fidelity FEM simulations."
        ]
    },
    "simulation_platform": {
        "title": "Automated Simulation Data Platform",
        "org": "UIUC",
        "date": "2026",
        "default_bullets": [
            "Built and executed an automated COMSOL-oriented simulation data platform generating 10,000 synthetic 2D polycrystalline microstructure samples across 228 controlled configurations with approximately 43-44 samples per configuration.",
            "Implemented Python, PowerShell, and Linux batch workflows for EBSD-style image generation, COMSOL material-image conversion, 10000 x 10000 label export, simulation execution, structured logging, and failure tracking.",
            "Refactored a COMSOL Java model to accept parameterized microstructure image inputs and assign material properties from grayscale phase labels."
        ]
    },
    "degradation_modeling": {
        "title": "Multiphysics Degradation Modeling Researcher",
        "org": "UIUC sponsored research project",
        "date": "2023-Present",
        "default_bullets": [
            "Developed EBSD-informed COMSOL/FEM corrosion models for aluminum alloys, linking microstructure, grain-boundary activity, and crystallographic orientation to localized degradation and material-loss evolution.",
            "Built image-processing workflows to quantify experimental corrosion/material-loss areas and compare them against FEM predictions for validation-oriented model refinement.",
            "Proposed a thin-film corrosion modeling roadmap extending SCD models toward TCD/Nernst-Planck transport with dynamic film thickness, oxygen limitation, and EBSD-informed kinetics."
        ]
    },
    "manufacturing_inspection": {
        "title": "Graduate Researcher, Manufacturing Quality AI Project",
        "org": "Foxconn Interconnect Technology / UIUC",
        "date": "2023-2025",
        "default_bullets": [
            "Built AOI pipelines for solder-joint localization, cropping, and defect classification across signal and ground joint types in industrial production images, improving transfer-learning test accuracy from roughly 86.0-91.2% to 92.5-94.0%.",
            "Improved binary defect detection through transfer learning and multi-view feature fusion, reaching up to 96.4% binary accuracy and 96.8% 5-fold binary accuracy with feature fusion.",
            "Validated new production samples with approximately 94.9% 3-fold accuracy, 96.0% precision, and 98.0% recall; integrated IRT temperature-curve signals to reach 100% bad-joint detection with about 6% over-kill.",
            "Delivered confidential performance metrics to the engineering team for downstream optimization while keeping production data private.",
            "Developed end-to-end deep learning models for ultrasonic metal welding quality prediction by fusing real-time sensor signals with process settings, improving prediction performance by at least 6% over conventional ML baselines."
        ]
    },
    "battery_thermal": {
        "title": "Battery Thermal Surrogate Modeling",
        "org": "UIUC",
        "date": "2024-2026",
        "default_bullets": [
            "Built and evaluated physics-informed CNN surrogate models for battery-pack heat-map prediction on an existing dataset, targeting fast thermal response estimation under varying boundary conditions.",
            "Investigated multi-fidelity and local-to-global surrogate strategies to reduce dependence on expensive high-fidelity thermal simulations."
        ]
    }
}

def match_project_key(title_str: str) -> str | None:
    """Helper to match LLM output project title to database key using simple keyword heuristics."""
    title_lower = title_str.lower()
    if "surrogate" in title_lower or "physics-guided" in title_lower:
        return "surrogate_modeling"
    elif "comsol" in title_lower or "platform" in title_lower or "polycrystalline" in title_lower:
        return "simulation_platform"
    elif "degradation" in title_lower or "corrosion" in title_lower or "multiphysics" in title_lower:
        return "degradation_modeling"
    elif "manufacturing" in title_lower or "welding" in title_lower or "foxconn" in title_lower or "inspection" in title_lower:
        return "manufacturing_inspection"
    elif "battery" in title_lower or "thermal" in title_lower:
        return "battery_thermal"
    return None

def build_tailored_docx(tailored_data: dict, output_path: Path):
    """Compiles the dynamically tailored resume DOCX file based on AI matching."""
    doc = Document()
    set_doc_defaults(doc)
    
    # Header
    add_header(
        doc,
        "Yulun Wu",
        "Scientific ML | Computational Materials | Multiphysics Simulation | Manufacturing AI"
    )
    
    # Dynamic Summary / Profile
    add_section(doc, "Summary")
    profile_text = tailored_data.get("tailored_summary", "").strip()
    if not profile_text:
        profile_text = "Ph.D. candidate combining scientific ML, multiphysics simulation, and manufacturing quality analytics for engineering decision support."
    add_inline_kv(doc, "Profile", profile_text)
    
    # Technical Skills (Keep stable/factual)
    add_section(doc, "Technical Skills")
    add_inline_kv(
        doc,
        "Simulation & Mechanics",
        "COMSOL Multiphysics, Abaqus, FEniCS, FEM/FEA, Level Set Methods, SolidWorks, nonlinear deformation, thermal-electrochemical multiphysics"
    )
    add_inline_kv(
        doc,
        "ML & Data",
        "Python, PyTorch, MATLAB, CNN, PINN, multi-task learning, continual learning, transfer learning, data imbalance mitigation, DOE"
    )
    add_inline_kv(
        doc,
        "Materials & Manufacturing",
        "Microstructure representation, EBSD-informed modeling, corrosion/degradation modeling, battery thermal maps, AOI inspection, sensor fusion"
    )
    
    # Selected Experience (Dynamically selected and tailored bullets)
    add_section(doc, "Selected Experience")
    selected_projects = tailored_data.get("selected_projects", [])
    
    used_keys = set()
    if selected_projects:
        for proj in selected_projects:
            title = proj.get("title", "")
            bullets = proj.get("bullets", [])
            
            db_key = match_project_key(title)
            if db_key and db_key not in used_keys:
                used_keys.add(db_key)
                db_proj = PROJECTS_DATABASE[db_key]
                add_role(doc, db_proj["title"], db_proj["org"], db_proj["date"])
                
                # Use tailored bullets if generated, otherwise default to inventory bullets
                bullets_to_use = bullets if bullets else db_proj["default_bullets"]
                for b in bullets_to_use:
                    add_bullet(doc, b)
                    
    # Fallback to default experience if LLM returned nothing or matching failed
    if not used_keys:
        fallback_keys = ["surrogate_modeling", "simulation_platform", "degradation_modeling"]
        for db_key in fallback_keys:
            db_proj = PROJECTS_DATABASE[db_key]
            add_role(doc, db_proj["title"], db_proj["org"], db_proj["date"])
            for b in db_proj["default_bullets"]:
                add_bullet(doc, b)
                
    # Education & Publications (Keep stable/factual)
    add_education(doc)
    add_publications(
        doc,
        [
            "Y. Wu, P. Bansal, Y. Li, Physics-informed multitask learning for microstructure-aware material metamodels of heterogeneous materials, Computational Materials Science, 2025.",
            "Y. Wu, Y. Li, Fusing Imbalanced Data via Physical Condition-Aware Surrogate Modeling, AIAA SciTech Forum, 2025.",
            "Y. Wu, Y. Meng, C. Shao, End-to-end online quality prediction for ultrasonic metal welding using sensor fusion and deep learning, Journal of Manufacturing Processes, 2022.",
            "Y. Cao, D. Guan, Y. Wu, et al., Box-level segmentation supervised deep neural networks for accurate and real-time multispectral pedestrian detection, ISPRS Journal, 2019."
        ]
    )
    
    doc.save(output_path)
    print_lg(f"Saved custom tailored resume to: {output_path}")

def generate_tailored_resume(job_description: str, job_id: str, company_name: str, model) -> Path | None:
    """API wrapper to match JD, query Gemini, and compile tailored resume docx."""
    from modules.ai.geminiConnections import gemini_tailor_resume_bullets
    
    # 1. Resolve Project Inventory Path
    project_inventory_path = os.environ.get("CAREER_PROJECT_INVENTORY", "career_project_inventory.md")
    if not os.path.exists(project_inventory_path) and os.path.exists("yulun_project_experience_inventory.md"):
        project_inventory_path = "yulun_project_experience_inventory.md"
        
    try:
        with open(project_inventory_path, "r", encoding="utf-8") as f:
            inventory_content = f.read()
    except Exception as e:
        critical_error_log(f"Failed to read project inventory file: {project_inventory_path}", e)
        return None
        
    # 2. Query Gemini for tailored summary and bullets
    try:
        print_lg(f"Requesting resume bullet tailoring from Gemini for Job ID: {job_id}...")
        tailored_data = gemini_tailor_resume_bullets(model, job_description, inventory_content)
        if not tailored_data or "error" in tailored_data:
            raise ValueError(f"Gemini API returned error: {tailored_data}")
    except Exception as e:
        critical_error_log("Failed to get tailored resume content from Gemini!", e)
        return None
        
    # 3. Compile and save the tailored DOCX file
    clean_company = re.sub(r'[^a-zA-Z0-9]', '_', company_name)
    filename = f"Yulun_Wu_Resume_{clean_company}_{job_id}.docx"
    output_dir = Path("all resumes")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / filename
    
    try:
        build_tailored_docx(tailored_data, output_path)
        return output_path
    except Exception as e:
        critical_error_log(f"Failed to compile tailored docx for {company_name}", e)
        return None
